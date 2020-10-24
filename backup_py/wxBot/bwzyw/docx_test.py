from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT
# 设置超链接
def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

#打开文档
document = Document()

#加入不同等级的标题
document.add_heading('Document Title',0)
document.add_heading(u'二级标题',1)
document.add_heading(u'二级标题',2)



#添加文本
paragraph = document.add_paragraph(u'添加了文本')

#add a hyperlink with the normal formatting (blue underline)
hyperlink = add_hyperlink(paragraph, 'http://www.google.com', 'Google', None, True)

#add a hyperlink with a custom color and no underline
hyperlink1 = add_hyperlink(paragraph, 'http://www.google.com', 'Google', 'FF8822', False)

#设置字号
# run = paragraph.add_run(u'设置字号')
run = paragraph.add_run(u'设置字号')
run.font.size=Pt(24)

#设置字体
run = paragraph.add_run('Set Font,')
run.font.name='Consolas'
#设置中文字体
run = paragraph.add_run(u'设置中\n文字体，')
run.font.name=u'宋体'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

#设置斜体
run = paragraph.add_run(u'斜体、')
run.italic = True

#设置粗体
run = paragraph.add_run(u'粗体').bold = True

#增加引用
document.add_paragraph('Intense quote', style='Intense Quote')

#增加有序列表
document.add_paragraph(
    u'有序列表元素1',style='List Number'
)
document.add_paragraph(
    u'有序列别元素2',style='List Number'
)

#增加无序列表
document.add_paragraph(
    u'无序列表元素1',style='List Bullet'
)
document.add_paragraph(
    u'无序列表元素2',style='List Bullet'
)

#增加图片（此处使用相对位置）
document.add_picture('1.jpg',width=Inches(1.25))

#增加表格
table = document.add_table(rows=3,cols=3)
hdr_cells=table.rows[0].cells
hdr_cells[0].text="第一列"
hdr_cells[1].text="第二列"
hdr_cells[2].text="第三列"

hdr_cells = table.rows[1].cells
hdr_cells[0].text = '2'
hdr_cells[1].text = 'aerszvfdgx'
hdr_cells[2].text = 'abdzfgxfdf'

hdr_cells = table.rows[2].cells
hdr_cells[0].text = '3'
hdr_cells[1].text = 'cafdwvaef'
hdr_cells[2].text = 'aabs zfgf'

#增加分页
document.add_page_break()
#保存文件

document.save('demo.docx')

