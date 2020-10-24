import requests
from pyquery import PyQuery as pq
import json
import traceback
import pdfkit
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
# 设置超链接
def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def get_mov(mov_name):
    headers = {
        'User-Agent': 'Mozilla / 5.0(Windows NT 10.0; Win64;x64) AppleWebKit / 537.36(KHTML, likeGecko) Chrome / 68.0.3440.106Safari / 537.36',
    }
    url = 'https://www.baiwanzy.com/index.php?m=vod-search&wd={}'.format(mov_name)
    try:
        r = requests.get(url, headers=headers)
        r.encoding = 'utf-8'
        html = pq(r.text)
        num = int(html('.nvc dl dd span:last-child').text())
        if num != 0:
            results = html('.xing_vb ul li').items()
            # print(list(results))
            # print(len(results))
            for r in results:
                print(r)
        # return num
    except:
        print(traceback.format_exc())
        return 0
# get_mov('庆')
# get_mov('庆余年')
def create_mov_link(url):
    mov_info = {}
    headers = {
        'User-Agent': 'Mozilla / 5.0(Windows NT 10.0; Win64;x64) AppleWebKit / 537.36(KHTML, likeGecko) Chrome / 68.0.3440.106Safari / 537.36',
    }
    try:
        r = requests.get(url, headers=headers,verify=False)
        r.encoding = 'utf-8'
        html = pq(r.text)
        vodTitle = html('.vodh').text()
        mov_info['title'] = ' '.join(vodTitle.split('\n'))
        vodInfos = list(html('.vodinfobox ul li').items())
        infos = []
        for v in vodInfos:
            infos.append(v.text())
        if infos:
            del infos[-3:]
        mov_info['info'] = infos
        vodLinks = html('.vodplayinfo:last-child div:nth-child(1) ul').items()
        values = list(vodLinks)[0]('li').items()
        links = []
        for v in values:
            links.append(v.text())
        if links:
            mov_info['link'] = links
        return mov_info
    except:
        print(traceback.format_exc())
        return 0
# def create_pdf():
#     mov_dict = create_mov_link('https://www.baiwanzy.com/?m=vod-detail-id-30207.html')
#     html = '<html><head><meta charset="UTF-8"></head><body><h1>{}</h1><br/><br/>{}<br/><br/>{}<br/><br/</body></html>'
#     mov_info = ''
#     for m in mov_dict['info']:
#         s = "{}<br/>".format(m)
#         mov_info += s
#     mov_link = ''
#     for i, m in enumerate(mov_dict['link']):
#         ms = m.split('$')
#         s = '{}  <a href="{}">播放</a>  '.format(ms[0], ms[1])
#         if (i+1) % 3 == 0:
#             s += '<br/>'
#         mov_link += s
#     config = pdfkit.configuration(wkhtmltopdf=r"wkhtmltopdf.exe")
#     # pdfkit.from_url('http://www.baidu.com', 'url_test.pdf',configuration=config) #通过url地址生成
#     pdfkit.from_string(html.format(mov_dict['title'], mov_info, mov_link), '庆余年.pdf', configuration=config)

def create_docx():
    mov_dict = create_mov_link('https://www.baiwanzy.com/?m=vod-detail-id-30207.html')
    # 打开文档
    document = Document()

    # 加入不同等级的标题
    document.add_heading(mov_dict['title'], 0)
    # 添加文本
    paragraph = document.add_paragraph()
    mov_info = ''
    for m in mov_dict['info']:
        s = "{}\n".format(m)
        mov_info += s
    # 设置中文字体
    run = paragraph.add_run(u'{}'.format(mov_info))
    run.font.name = u'微软雅黑'
    run.font.size = Pt(14)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')



    for i, m in enumerate(mov_dict['link']):
        ms = m.split('$')
        # 设置中文字体
        if i == 0:
            tstr = '\n\n{}（播放地址已被WX屏蔽，请在浏览器中打开）\n'.format(ms[0])
        else:
            tstr = '\n{}（播放地址已被WX屏蔽，请在浏览器中打开）\n'.format(ms[0])
        run = paragraph.add_run(u'{}'.format(tstr))
        run.font.name = u'微软雅黑'
        run.font.size = Pt(14)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        # if (i+1) % 3 == 0:
        #     sstr = '播放\n'
        # else:
        #     sstr = '播放  '

        add_hyperlink(paragraph, ms[1], '{}\n'.format(ms[1]), 'B22222', False)

    # config = pdfkit.configuration(wkhtmltopdf=r"wkhtmltopdf.exe")
    # # pdfkit.from_url('http://www.baidu.com', 'url_test.pdf',configuration=config) #通过url地址生成
    # pdfkit.from_string(html.format(mov_dict['title'], mov_info, mov_link), '庆余年.pdf', configuration=config)

    document.save('demo.docx')
def yz_html(i):
    headers = {
        'User-Agent': 'Mozilla / 5.0(Windows NT 10.0; Win64;x64) AppleWebKit / 537.36(KHTML, likeGecko) Chrome / 68.0.3440.106Safari / 537.36',
    }
    url = 'http://dcs.yozosoft.com/onlinefile'
    dict = {
        'downloadUrl': 'http://ele.379lb.cn/wenku/{}.doc'.format(i),
        'convertType': 1,
        'isAsync': 0,
        'sourceFolder': '771906ca-68e9-4aed-8529-104aaba22b68',
        'isDownload': 0,
        'isSignature': 0
    }
    try:
        r = requests.post(url, headers=headers, data=dict)
        print(r.text)
    except:
        print(traceback.format_exc())
        return 0
#
# for i in range(1):
#     yz_html(i+1)
# create_pdf()
# create_docx()
def short_url(long_url): #将长链接变成短链接
    host = 'https://dwz.cn'
    path = '/admin/v2/create'
    url = host + path
    url1 = long_url
    content_type = 'application/json'
    # TODO: 设置Token
    token = '9b2b162c0770cbba0da6c1cbcee36f54'  #访问"https://dwz.cn/console/userinfo"登录百度账号，即可免费获取令牌
    # TODO：设置待创建的长网址
    bodys = {'url': url1, 'TermOfValidity': '1-year'}
    # 配置headers
    headers = {'Content-Type': content_type, 'Token': token}
    # 发起请求
    response = requests.post(url=url, data=json.dumps(bodys), headers=headers).json()
    # 读取响应
    if response['Code'] == 0:
        return response["ShortUrl"]
    else:
        return url1
print(short_url('https://mp.weixin.qq.com/s/6xIbrbkHZ0QiwDHybIUetg'))